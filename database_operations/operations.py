import shutil
import sqlite3
import csv
from datetime import datetime
from docx import Document
import os
import matplotlib.pyplot as plt
import seaborn as sns
sns.set()


class Report:
    def __init__(self, output, b_categories, a_categories=None, number=5):
        self.current_datetime = self.return_date_time('%Y-%m-%d %H:%M:%S')
        self.current_date = self.return_date_time('%Y_%m_%d')
        self.number = number
        self.document = Document()
        self.output = output
        self.columns = {'test': (),
                        'work': ('Work',),
                        'total': ('Work', 'Total')}
        self.basic_categories = b_categories
        self.advanced_categories = a_categories

    def add_basic_content(self):
        self.document.add_heading('Report', 0)
        section = self.document.sections[0]
        header = section.header

        paragraph = header.paragraphs[0]
        paragraph.text = f"\t\t{self.current_datetime}"
        paragraph.style = self.document.styles["Header"]

        self.document.add_heading('Best students', level=1)

        for query in self.basic_categories:
            paragraph = self.document.add_paragraph()
            paragraph.style = 'List Bullet'
            run = paragraph.add_run(f'Best {query}')
            run.bold = True

            table = self.document.add_table(
                rows=1, cols=4+len(self.columns[query]))
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Place'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'City'
            hdr_cells[3].text = 'Test'
            if self.columns[query]:
                for i, col in enumerate(self.columns[query]):
                    hdr_cells[4+i].text = col

            records = best_results(query=query, number=self.number, raw=True)

            counter = 1
            for items in records:
                row_cells = table.add_row().cells
                row_cells[0].text = str(counter)
                counter += 1
                row_cells[1].text = f'{items[2]} {items[3]}'
                row_cells[2].text = items[4]
                row_cells[3].text = str(items[5])
                if query == 'work' or query == 'total':
                    row_cells[4].text = str(items[6])
                    if query == 'total':
                        row_cells[5].text = str(int(items[5]) + int(items[6]))
            

    def add_advanced_content(self):
        self.document.add_heading('Advanced analysis', level=1)
        
        mode_translation = {1: 'Sex comparison',
                            2: 'City comparison'}
        
        if self.advanced_categories is not None:
            for i in self.advanced_categories:
                paragraph = self.document.add_paragraph()
                paragraph.style = 'List Bullet'
                run = paragraph.add_run(mode_translation[i])
                run.bold = True
                create_graphs(mode=i)
                for i in range(len(os.listdir('graphs'))):
                    self.document.add_picture(f'graphs/{i}.png')
                shutil.rmtree('graphs')

    def save(self):
        self.document.save(f'{self.current_date}_{self.output}')

    @staticmethod
    def return_date_time(format):
        return datetime.now().strftime(format)


def create_graphs(mode=1):
    if not os.path.exists('graphs'):
        os.mkdir('graphs')

    results = advanced(mode=mode)

    properties = {1: ['Students number', 'pie', None],
                  2: ['Average test points', 'bar', 'Points'],
                  3: ['Average work points', 'bar', 'Points'],
                  4: ['Average total points', 'bar', 'Points']}


    for i in range(len(results)):
        if i == 0:
            continue

        values = [item[1] for item in results[i]]
        names1 = [item[0] for item in results[i]]

        if properties[i][1] == 'pie':
            plt.pie(values, labels=names1, autopct=make_autopct(values))
        else:
            plt.bar(names1, values)

        plt.title(properties[i][0])
        plt.ylabel(properties[i][2])

        plt.savefig(f'graphs/{i-1}.png')
        plt.close()


def make_autopct(values):
    def my_autopct(pct):
        total = sum(values)
        val = int(round(pct*total/100.0))
        return '{p:.1f}%  ({v:d})'.format(p=pct, v=val)
    return my_autopct


def advanced(path='students.db', mode=1):
    conn = sqlite3.connect(path)
    cursor = conn.cursor()
    results = []
    if mode == 1:
        query = 'sex'
    else:
        query = 'city'

    cursor.execute(f'''SELECT
                    COUNT(id)
                    FROM
                    STUDENTS''')

    results.append(cursor.fetchall())

    cursor.execute(f'''SELECT {query},
                    COUNT(*)
                    AS
                    `number`
                    FROM
                    STUDENTS
                    GROUP
                    BY
                    {query}''')

    results.append(cursor.fetchall())

    cursor.execute(f'''SELECT {query},
                    AVG(test)
                    AS
                    `test`
                    FROM
                    STUDENTS
                    GROUP
                    BY
                    {query}''')

    records = [(x, round(y, 2)) for x, y in cursor.fetchall()]
    results.append(records)

    cursor.execute(f'''SELECT {query},
                    AVG(work)
                    AS
                    `work`
                    FROM
                    STUDENTS
                    GROUP
                    BY
                    {query}''')

    records = [(x, round(y, 2)) for x, y in cursor.fetchall()]
    results.append(records)

    cursor.execute(f'''SELECT {query},
                    AVG(test + work)
                    AS
                    `total`
                    FROM
                    STUDENTS
                    GROUP
                    BY
                    {query}''')

    records = [(x, round(y, 2)) for x, y in cursor.fetchall()]
    results.append(records)

    if query == 'sex':
        for records in results:
            for idx, items in enumerate(records):
                if len(items) == 1:
                    continue
                elif items[0] == 'm':
                    records[idx] = ('Men', items[1])
                else:
                    records[idx] = ('Women', items[1])

    conn.close()
    return results


def best_results(query, number, raw=False, path='students.db'):

    queries = {'test': 'cast(test as int)',
               'work': 'cast(work as int)',
               'total': 'cast(test as int) + cast(work as int)'}

    conn = sqlite3.connect(path)
    cursor = conn.cursor()
    students_info = []

    cursor.execute(F'''SELECT * FROM STUDENTS
        ORDER BY {queries[query]} DESC
        LIMIT {number}''')

    if raw:
        return cursor.fetchall()

    for place, student in enumerate(cursor.fetchall()):
        tmp = []
        place = f'*{place + 1} place'
        tmp.append(f'\tFirst Name: {student[2]}')
        tmp.append(f'\tLast Name: {student[3]}')
        tmp.append(f'\tCity: {student[4]}')
        tmp.append(f'\tTest points: {student[5]}')
        if query == 'total':
            tmp.append(f'\tTOTAL: {int(student[5]) + int(student[6])}')
        students_info.append([place, tmp])
    conn.close()
    return students_info


def save_database(db_file, students_list):
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()
    cursor.execute('DROP TABLE STUDENTS')
    conn.execute('''CREATE TABLE STUDENTS
    (id, sex, first_name, last_name, city, test, work)''')

    for idx, student in enumerate(students_list):
        info = list(student.save_mode())
        info.insert(0, idx + 1)
        info = tuple(info)
        cursor.execute(f'''INSERT INTO STUDENTS
        (id, sex, first_name, last_name, city, test, work)
        VALUES {info}''')

    print('Database updated successfully!')
    conn.commit()
    conn.close()


def create_from_csv(path):
    """Create database from csv file

    Args:
        path (str): Path to csv file
    """
    conn = sqlite3.connect('students.db')
    conn.execute('''CREATE TABLE STUDENTS 
    (id, sex, first_name, last_name, city, test, work)''')
    cursor = conn.cursor()

    with open(path, 'r') as csv_file:
        reader = csv.reader(csv_file)
        for idx, row in enumerate(reader):
            row.insert(0, idx + 1)
            row = tuple(row)
            cursor.execute(f'''INSERT INTO STUDENTS 
            (id, sex, first_name, last_name, city, test, work)
            VALUES {row}''')

    conn.commit()
    conn.close()


if __name__ == '__main__':
    report = Report('report.docx', ['test', 'work', 'total'])
    report.add_basic_content()
    report.add_advanced_content()
    report.save()
    # create_graphs()
